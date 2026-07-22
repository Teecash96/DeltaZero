from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "deliverables/DeltaZero_Demo_Video_Script_v1.docx"
GREEN = "B7FF4A"
DARK = "102018"
MUTED = "66736B"
LIGHT = "EEF5F0"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_title(doc, text, size=28, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("527A26")
    r.font.letter_spacing = Pt(1)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    cell = table.cell(0, 0)
    cell.width = Inches(6.35)
    margins(cell, 150, 180, 150, 180)
    shade(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("527A26")
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    return table


def add_shot_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.55, 0.78, 2.25, 2.77]
    headers = ["Shot", "Time", "Screen action", "Voiceover"]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(w)
        shade(cell, DARK)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    for idx, timing, action, voice in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        values = [idx, timing, action, voice]
        for i, (value, width) in enumerate(zip(values, widths)):
            cells[i].width = Inches(width)
            margins(cells[i], 100, 120, 100, 120)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(value)
            r.font.name = "Arial"
            r.font.size = Pt(8.5)
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string("527A26")
        if len(table.rows) % 2 == 1:
            for c in cells:
                shade(c, "F7FAF8")
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18
for level, size in ((1, 17), (2, 13), (3, 11)):
    s = styles[f"Heading {level}"]
    s.font.name = "Arial"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(DARK if level == 1 else "355143")
    s.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    s.paragraph_format.space_after = Pt(6)
for list_style in ("List Bullet", "List Number"):
    s = styles[list_style]
    s.font.name = "Arial"
    s.font.size = Pt(10.5)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.15

header = section.header
p = header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("DELTAZERO  /  DEMO PRODUCTION GUIDE")
r.font.name = "Arial"
r.font.size = Pt(8)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(MUTED)
footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DeltaZero - Know your hedge. Protect your capital.")
r.font.name = "Arial"
r.font.size = Pt(8)
r.font.color.rgb = RGBColor.from_string(MUTED)

add_kicker(doc, "Editor-ready run of show")
add_title(doc, "DeltaZero Demo Video Script")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run("A 100-second product walkthrough for the OKX.AI hackathon")
r.font.name = "Arial"
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string(MUTED)

meta = doc.add_table(rows=2, cols=3)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
meta_values = [
    ("Target length", "90-110 seconds"),
    ("Format", "16:9, 1080p MP4"),
    ("Tone", "Institutional, calm, precise"),
    ("Primary URL", "delta-zero-alpha.vercel.app"),
    ("Live price", "1 USDT per paid call"),
    ("Agent", "OKX Agent #5739"),
]
for i, (label, value) in enumerate(meta_values):
    cell = meta.cell(i // 3, i % 3)
    cell.width = Inches(2.12)
    margins(cell, 120, 140, 120, 140)
    shade(cell, LIGHT if i % 2 == 0 else "F7FAF8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    rr = p.add_run(label.upper())
    rr.font.size = Pt(7.5)
    rr.font.bold = True
    rr.font.color.rgb = RGBColor.from_string("527A26")
    p2 = cell.add_paragraph(value)
    p2.paragraph_format.space_after = Pt(0)
    p2.runs[0].font.size = Pt(9.5)
    p2.runs[0].font.bold = True

doc.add_paragraph()
add_callout(doc, "Core story", "DeltaZero converts complex pseudo-delta-neutral risk into deterministic operator decisions, then proves monetization through a real OKX x402 payment and an independently verifiable on-chain receipt.")

add_heading(doc, "What the finished video must prove")
for text in [
    "The product works end to end, not just as a collection of screens.",
    "The free Judge Demo explains value before asking the reviewer to fund a wallet.",
    "The live Risk Engine protects paid calculations behind a 1 USDT OKX x402 boundary.",
    "A completed payment produces a real X Layer transaction hash and receipt evidence.",
    "DeltaZero is read-only decision support: no custody, approvals, signatures, or automatic trade execution.",
]:
    add_bullet(doc, text)

add_heading(doc, "Recording preparation")
for text in [
    "Use Chrome at 1440 x 900 or 1920 x 1080. Set browser zoom to 100% and use dark mode.",
    "Close unrelated tabs, notifications, password managers, email, developer consoles, and terminal windows.",
    "Never show API credentials, admin keys, seed phrases, private wallet information, Railway variables, or browser autofill suggestions.",
    "Prepare an OKX Wallet funded on X Layer with at least 1 USD₮0 plus enough OKB for gas.",
    "Open these tabs in advance: Homepage, Judge Demo, Risk Engine, Agent Console, Hyperliquid Live, and the X Layer transaction explorer.",
    "Use one stable reference scenario: SOL, 5,000 USD capital, medium risk, 14% long yield, 3% short funding, and 1% fee drag.",
]:
    add_bullet(doc, text)

doc.add_page_break()
add_kicker(doc, "Run of show")
add_title(doc, "Shot-by-Shot Script", 23)

shots = [
    ("01", "0:00-0:06", "Homepage. Hold on the headline, then move toward Judge Demo.", "DeltaZero is deterministic DeFi risk intelligence for agents and users running pseudo-delta-neutral strategies."),
    ("02", "0:06-0:16", "Open Judge Demo. Show the four reference workflows and one verdict. Keep the 'No payment required' label visible.", "The free Judge Demo explains the workflow through verified reference scenarios: strategy construction, hedge-drift auditing, funding stress testing, and Monte Carlo sensitivity."),
    ("03", "0:16-0:24", "Return to the homepage and open Launch Risk Engine. Show all four service cards.", "For live inputs, DeltaZero's Risk Engine applies transparent formulas and produces a clear operator verdict, Safety Buffer, and risk zone."),
    ("04", "0:24-0:36", "Select Strategy Build. Enter the prepared SOL assumptions and submit. Pause on the 1 USDT payment screen.", "Live calculations are protected by OKX x402. One paid call costs one USDT and returns a verifiable payment challenge before analysis runs."),
    ("05", "0:36-0:50", "Connect OKX Wallet, review network/asset/receiver, approve payment, and wait for the result. Do not cut away during the confirmation state.", "The user reviews the X Layer payment, authorizes it in their own wallet, and the protected endpoint verifies settlement before returning the analysis."),
    ("06", "0:50-1:02", "Show the completed DeltaZero Verdict, allocation, net carry, hedge drift, Safety Buffer, and operator action.", "The result is deterministic decision support. DeltaZero does not invent financial numbers and never requests custody or trade execution."),
    ("07", "1:02-1:13", "Open the On-chain Payment Receipt panel. Show status, transaction hash, payer, receiver, amount, asset, block, and verified transfer event.", "Every successful payment produces evidence: the transaction hash, payer, receiver, asset, amount, block number, and verified token transfer event."),
    ("08", "1:13-1:21", "Click View on OKLink, then return. Briefly show Copy receipt and Download JSON.", "The receipt can be checked independently on the X Layer explorer or exported as raw JSON for audit and automation."),
    ("09", "1:21-1:29", "Open Agent Console. Show the monitoring state and deterministic execution payload, without claiming a real trade executed.", "Agents can monitor drift and consume structured recommendations through the API and MCP interfaces. Execution remains controlled by the operator."),
    ("10", "1:29-1:37", "Open Hyperliquid Live and briefly show free market context. Then return to homepage.", "Free Hyperliquid context complements the premium Risk Engine, helping users inspect current conditions before running paid analysis."),
    ("11", "1:37-1:41", "End on the homepage headline and logo.", "DeltaZero. Know your hedge. Protect your capital."),
]
add_shot_table(doc, shots)

doc.add_page_break()
add_kicker(doc, "Narration")
add_title(doc, "Clean Voiceover Script", 23)
voice = [
    "DeltaZero is deterministic DeFi risk intelligence for agents and users running pseudo-delta-neutral strategies.",
    "The free Judge Demo explains the workflow through verified reference scenarios: strategy construction, hedge-drift auditing, funding stress testing, and Monte Carlo sensitivity.",
    "For live inputs, DeltaZero's Risk Engine applies transparent formulas and produces a clear operator verdict, Safety Buffer, and risk zone.",
    "Live calculations are protected by OKX x402. One paid call costs one USDT and returns a verifiable payment challenge before analysis runs.",
    "The user reviews the X Layer payment, authorizes it in their own wallet, and the protected endpoint verifies settlement before returning the analysis.",
    "The result is deterministic decision support. DeltaZero does not invent financial numbers and never requests custody or trade execution.",
    "Every successful payment produces evidence: the transaction hash, payer, receiver, asset, amount, block number, and verified token transfer event.",
    "The receipt can be checked independently on the X Layer explorer or exported as raw JSON for audit and automation.",
    "Agents can monitor drift and consume structured recommendations through the API and MCP interfaces. Execution remains controlled by the operator.",
    "Free Hyperliquid context complements the premium Risk Engine, helping users inspect current conditions before running paid analysis.",
    "DeltaZero. Know your hedge. Protect your capital.",
]
for paragraph in voice:
    p = doc.add_paragraph(paragraph)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25

add_callout(doc, "Narration rule", "Use a confident, measured voice at approximately 135-145 words per minute. Avoid hype, exaggerated claims, and language suggesting guaranteed profitability or autonomous trade execution.")

on_screen_heading = add_heading(doc, "On-screen labels")
on_screen_heading.paragraph_format.page_break_before = True
for text in [
    "FREE JUDGE PREVIEW - VERIFIED REFERENCE SCENARIOS",
    "LIVE RISK ENGINE - 1 USDT VIA OKX x402",
    "DETERMINISTIC CALCULATIONS - NO LLM-GENERATED NUMBERS",
    "READ ONLY - NO CUSTODY - NO TRADE EXECUTION",
    "VERIFIABLE X LAYER PAYMENT RECEIPT",
    "API + MCP READY FOR AGENT WORKFLOWS",
]:
    add_bullet(doc, text)

add_heading(doc, "Editing direction")
for text in [
    "Use straight cuts or subtle 150-250 ms dissolves. No dramatic zooms, glitch effects, neon overlays, or fast cursor circles.",
    "Keep background music optional and below -24 LUFS so narration remains dominant.",
    "Add captions in a clean sans-serif font. Use white text with restrained lime emphasis for product terms only.",
    "Zoom only when needed to make payment and receipt fields legible. Keep the real transaction hash visible long enough to pause and verify.",
    "Do not fabricate a settled payment. If settlement has not succeeded, postpone the final recording rather than replacing it with a mock receipt.",
]:
    add_bullet(doc, text)

doc.add_page_break()
add_kicker(doc, "Quality gate")
add_title(doc, "Final Recording Checklist", 23)

sections = {
    "Product": [
        "Production frontend and backend are both healthy.",
        "Judge Demo is accessible without payment and is clearly labelled as reference output.",
        "Risk Engine requests exactly 1 USDT on X Layer.",
        "The paid result unlocks after wallet settlement.",
        "The receipt panel shows a real transaction hash and verified transfer event.",
    ],
    "Evidence": [
        "Transaction opens successfully on OKLink.",
        "Receiver matches 0x2cc85ed0c35fc2e4e7ad28592a52a4cda96b44da.",
        "Asset and amount match the payment challenge.",
        "Receipt JSON downloads and contains the settlement evidence.",
        "No admin or demo-access bypass is used in the paid segment.",
    ],
    "Safety and accuracy": [
        "No private keys, credentials, wallet recovery phrases, or personal notifications appear.",
        "No claim of guaranteed profit, prediction accuracy, custody, or automatic execution.",
        "Judge Demo values are described as reference scenarios, not live customer results.",
        "Agent Console is described as monitoring and structured recommendation output.",
    ],
    "Export": [
        "Final duration is between 90 and 110 seconds.",
        "Export is 1920 x 1080, 16:9, H.264 MP4.",
        "Text is readable on a phone at normal playback size.",
        "Audio is normalized, captions are checked, and no UI is cropped.",
    ],
}
for heading, items in sections.items():
    checklist_heading = add_heading(doc, heading, 2)
    if heading == "Safety and accuracy":
        checklist_heading.paragraph_format.page_break_before = True
    for item in items:
        add_bullet(doc, "[ ] " + item)

add_heading(doc, "Release decision", 2)
p = doc.add_paragraph("Do not publish until the real 1 USDT settlement and downloadable receipt are verified in production.")
p.paragraph_format.keep_together = True

doc.core_properties.title = "DeltaZero Demo Video Script"
doc.core_properties.subject = "OKX.AI hackathon product demo production guide"
doc.core_properties.author = "Akanbi Labs"
doc.core_properties.keywords = "DeltaZero, OKX, x402, DeFi, demo script"
doc.save(OUT)
print(OUT)
